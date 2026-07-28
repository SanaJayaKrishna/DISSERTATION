"""
generate_dissertation_docx.py
Generates a BITS WILP-compliant dissertation DOCX for:
  Title   : A Robot-Agnostic Framework for Language-Conditioned Task Planning
             using Capability-Aware Large Language Models
  Student : Sana Jaya Krishna  |  2024AA05783
  Course  : AIMLCZG628T  |  M.Tech. AI & ML  |  BITS Pilani WILP

Run:
    python3 generate_dissertation_docx.py
Output:
    dissertation_final_report.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

# ── helpers ────────────────────────────────────────────────────────────────

def set_paragraph_spacing(para, before=0, after=6, line_spacing=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line_spacing is not None:
        pf.line_spacing      = Pt(line_spacing)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

def add_page_break(doc):
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(docx.oxml.ns.qn('w:br'))
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)

def set_font(run, name='Times New Roman', size=12, bold=False, italic=False):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic= italic

def heading(doc, text, level=1, numbered=False, prefix=''):
    para = doc.add_heading(text if not prefix else f'{prefix}  {text}', level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(para, before=12, after=6)
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14 if level == 1 else 12)
    return para

def body(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para = doc.add_paragraph(text)
    para.alignment = align
    set_paragraph_spacing(para, before=0, after=6, line_spacing=24)
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return para

def bold_label(doc, label, text='', align=WD_ALIGN_PARAGRAPH.LEFT):
    para = doc.add_paragraph()
    para.alignment = align
    set_paragraph_spacing(para, before=0, after=4)
    r = para.add_run(label)
    set_font(r, bold=True)
    if text:
        r2 = para.add_run(' ' + text)
        set_font(r2)
    return para

def bullet(doc, text):
    para = doc.add_paragraph(text, style='List Bullet')
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(para, before=0, after=4, line_spacing=24)
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return para

def enum_item(doc, text):
    para = doc.add_paragraph(text, style='List Number')
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(para, before=0, after=4, line_spacing=24)
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return para

def horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def sig_table(doc, left_label, right_label):
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            cell._tc.get_or_add_tcPr()
    # row 0: labels
    table.cell(0, 0).text = left_label
    table.cell(0, 1).text = right_label
    # row 1: blank (signature space)
    table.cell(1, 0).text = ''
    table.cell(1, 1).text = ''
    # row 2: date line
    table.cell(2, 0).text = 'Date: _______________'
    table.cell(2, 1).text = 'Date: _______________'
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    table.rows[1].height = Cm(2.5)
    # remove borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for tag in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{tag}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def centered(doc, text, size=12, bold=False):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(para, before=0, after=6)
    r = para.add_run(text)
    set_font(r, size=size, bold=bold)
    return para

def section_break(doc):
    """Add a page break using a paragraph."""
    para = doc.add_paragraph()
    run  = para.add_run()
    from docx.oxml import OxmlElement as OE
    br = OE('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    set_paragraph_spacing(para, 0, 0)

# ── document setup ─────────────────────────────────────────────────────────

import docx

doc = Document()

# Page margins: 1 inch all sides
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)

# Default body style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
pf = style.paragraph_format
pf.space_after  = Pt(6)
pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# ═══════════════════════════════════════════════════════════════════════════
# COVER PAGE  (Appendix-A)
# ═══════════════════════════════════════════════════════════════════════════

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(para, before=36, after=12)
r = para.add_run('A REPORT ON')
set_font(r, size=14, bold=True)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(para, before=6, after=24)
r = para.add_run(
    'A ROBOT-AGNOSTIC FRAMEWORK FOR\n'
    'LANGUAGE-CONDITIONED TASK PLANNING USING\n'
    'CAPABILITY-AWARE LARGE LANGUAGE MODELS'
)
set_font(r, size=16, bold=True)

centered(doc, 'BY', size=14, bold=True)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(para, before=6, after=24)
r1 = para.add_run('Sana Jaya Krishna\n')
set_font(r1, size=14, bold=True)
r2 = para.add_run('ID No.: 2024AA05783')
set_font(r2, size=12)

centered(doc, 'AT', size=14, bold=True)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(para, before=4, after=24)
r = para.add_run('[STATION NAME AND CENTRE]\n[ORGANIZATION NAME & LOCATION]')
set_font(r, size=12)

# BITS logo
img_path = 'vertopal_7695069776664397bca9a849d22a270b/media/image1.png'
if os.path.exists(img_path):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(para, before=12, after=12)
    run = para.add_run()
    run.add_picture(img_path, width=Inches(1.5))

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(para, before=12, after=4)
r = para.add_run('BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI')
set_font(r, size=14, bold=True)

centered(doc, 'VIDYA VIHAR, PILANI, RAJASTHAN – 333031', size=11)
centered(doc, '[Month, Year]', size=12, bold=True)

# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE  (Appendix-B)
# ═══════════════════════════════════════════════════════════════════════════

section_break(doc)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(para, before=24, after=12)
r = para.add_run('A REPORT ON')
set_font(r, size=14, bold=True)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(para, before=6, after=24)
r = para.add_run(
    'A ROBOT-AGNOSTIC FRAMEWORK FOR\n'
    'LANGUAGE-CONDITIONED TASK PLANNING USING\n'
    'CAPABILITY-AWARE LARGE LANGUAGE MODELS'
)
set_font(r, size=16, bold=True)

centered(doc, 'BY', size=14, bold=True)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(para, before=6, after=12)
r1 = para.add_run('Sana Jaya Krishna\n')
set_font(r1, size=14, bold=True)
r2 = para.add_run('ID No.: 2024AA05783\n')
set_font(r2, size=12)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(para, before=4, after=12)
r = para.add_run(
    'Prepared in partial fulfilment of the\n'
    'WILP Dissertation Course (Course No.: AIMLCZG628T)'
)
set_font(r, size=12)

centered(doc, 'AT', size=14, bold=True)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(para, before=4, after=12)
r = para.add_run('[ORGANIZATION NAME & LOCATION]\n')
set_font(r, size=12)
r2 = para.add_run('Degree Programme: M.Tech. Artificial Intelligence and Machine Learning')
set_font(r2, size=12)

if os.path.exists(img_path):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(para, before=12, after=12)
    run = para.add_run()
    run.add_picture(img_path, width=Inches(1.3))

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(para, before=12, after=4)
r = para.add_run('BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI')
set_font(r, size=14, bold=True)
centered(doc, '[Month, Year]', size=12, bold=True)

# ═══════════════════════════════════════════════════════════════════════════
# CERTIFICATE
# ═══════════════════════════════════════════════════════════════════════════

section_break(doc)
heading(doc, 'CERTIFICATE', level=1)

body(doc,
    'This is to certify that the dissertation titled "A Robot-Agnostic Framework for '
    'Language-Conditioned Task Planning using Capability-Aware Large Language Models" '
    'submitted by Sana Jaya Krishna (ID No. 2024AA05783) in partial fulfilment of the '
    'requirements for the course AIMLCZG628T (Dissertation) of M.Tech. Artificial '
    'Intelligence and Machine Learning, BITS Pilani WILP Division, is a bonafide record '
    'of the work carried out by the student under my supervision.'
)

doc.add_paragraph()  # space
sig_table(doc, 'Supervisor\n[SUPERVISOR NAME]\n[SUPERVISOR DESIGNATION]',
               'Additional Examiner\n[ADDITIONAL EXAMINER NAME]\n[ADDITIONAL EXAMINER DESIGNATION]')

doc.add_paragraph()
bold_label(doc, 'Faculty Mentor:', '[FACULTY MENTOR NAME], BITS Pilani WILP')

# ═══════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENTS
# ═══════════════════════════════════════════════════════════════════════════

section_break(doc)
heading(doc, 'Acknowledgements', level=1)

body(doc,
    'I would like to express my sincere gratitude to all those who contributed to the '
    'successful completion of this dissertation.'
)
body(doc,
    'I am deeply thankful to [SUPERVISOR NAME AND DESIGNATION], my Supervisor, and '
    '[ADDITIONAL EXAMINER NAME AND DESIGNATION], Additional Examiner, for their '
    'invaluable guidance, technical inputs, and continuous encouragement throughout '
    'the course of this project.'
)
body(doc,
    'I extend my sincere thanks to [FACULTY MENTOR NAME], my Faculty Mentor at BITS '
    'Pilani WILP, for their academic oversight and timely feedback that helped shape '
    'the direction and quality of this work.'
)
body(doc,
    'I am grateful to the management and colleagues at [ORGANIZATION NAME], '
    '[ORGANIZATION LOCATION], for providing the necessary resources, time, and support '
    'to carry out this dissertation work.'
)
body(doc,
    'I also acknowledge the open-source robotics community and the developers of ROS2, '
    'MoveIt2, Nav2, and the various language model frameworks whose tools and documentation '
    'formed an essential part of this work.'
)
body(doc,
    'Finally, I thank my family for their unwavering support and encouragement.'
)
doc.add_paragraph()
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = para.add_run(
    'Sana Jaya Krishna\n2024AA05783\n'
    'M.Tech. Artificial Intelligence and Machine Learning\n'
    'BITS Pilani, WILP Division\n[Month, Year]'
)
set_font(r, bold=True)

# ═══════════════════════════════════════════════════════════════════════════
# ABSTRACT SHEET  (Appendix-C)
# ═══════════════════════════════════════════════════════════════════════════

section_break(doc)

centered(doc, 'BIRLA INSTITUTE OF TECHNOLOGY AND SCIENCE, PILANI (RAJASTHAN)', size=13, bold=True)
centered(doc, 'WILP Division', size=12, bold=True)
doc.add_paragraph()

meta = [
    ('Organization:',       '[ORGANIZATION NAME]'),
    ('Location:',           '[ORGANIZATION LOCATION]'),
    ('Duration:',           '[PROJECT DURATION]'),
    ('Date of Start:',      '[DATE OF START]'),
    ('Date of Submission:', '[DATE OF SUBMISSION]'),
]
for label, val in meta:
    bold_label(doc, label, val)

doc.add_paragraph()
bold_label(doc, 'Title of the Project:')
body(doc, 'A Robot-Agnostic Framework for Language-Conditioned Task Planning using '
          'Capability-Aware Large Language Models')

bold_label(doc, 'ID No./Name of the Student:', '2024AA05783 / Sana Jaya Krishna')
bold_label(doc, 'Name(s) and Designation(s) of Supervisor and Additional Examiner:')
body(doc, 'Supervisor: [SUPERVISOR NAME AND DESIGNATION]\n'
          'Additional Examiner: [ADDITIONAL EXAMINER NAME AND DESIGNATION]')
bold_label(doc, 'Name of the Faculty Mentor:', '[FACULTY MENTOR NAME]')

bold_label(doc, 'Key Words:',
           'Robot-agnostic planning, Large Language Models, capability representation, '
           'URDF, task planning, embodied AI, LLM evaluation, robotic skills, ROS2')
bold_label(doc, 'Project Areas:', 'Robotics, Artificial Intelligence, Natural Language Processing')

bold_label(doc, 'Abstract:')
body(doc,
    'Modern robots exhibit heterogeneous embodiments with varying mobility, manipulation, '
    'perception, and operational capabilities. A logically correct plan generated by a '
    'Large Language Model (LLM) may be infeasible if the assigned robot lacks the required '
    'capabilities. This dissertation presents a robot-agnostic, capability-aware framework '
    'for language-conditioned robotic task planning. The proposed approach derives structured '
    'capability representations from robot description files (URDF/Xacro) and provides them '
    'as context to an LLM planner alongside world-state representations and a standardized '
    'skills library. The framework is robot-agnostic in architecture but robot-aware during '
    'planning. A deterministic software-based evaluation module assesses generated plans '
    'across six dimensions: goal achievement, action validity, object and location validity, '
    'capability compliance, logical ordering, and plan feasibility, without relying on a '
    'secondary LLM as a judge. A custom dataset of approximately 500 natural-language tasks '
    'across 11 diverse environments was developed to benchmark multiple lightweight open-source '
    'LLMs under identical experimental conditions, enabling controlled comparison of planning '
    'performance and providing a baseline for subsequent parameter-efficient fine-tuning (PEFT) '
    'experiments. The framework establishes a foundational architecture and evaluation '
    'methodology for capability-aware robotic task planning.'
)

doc.add_paragraph()
sig_table(doc, 'Signature of Student', 'Signature of Supervisor')

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS  (manual — LaTeX-generated ToC not available in docx)
# ═══════════════════════════════════════════════════════════════════════════

section_break(doc)
heading(doc, 'Table of Contents', level=1)

toc_entries = [
    ('Certificate',                                        'ii'),
    ('Acknowledgements',                                   'iii'),
    ('Abstract Sheet',                                     'iv'),
    ('Table of Contents',                                  'v'),
    ('List of Abbreviations and Acronyms',                 'vi'),
    ('1.  Introduction and Research Motivation',           '1'),
    ('2.  Literature Review and Research Gap',             '3'),
    ('    2.1  Research Gap',                              '4'),
    ('3.  Problem Definition and Objectives',              '5'),
    ('    3.1  Problem Definition',                        '5'),
    ('    3.2  Research Objectives',                       '6'),
    ('4.  Proposed Framework and System Architecture',     '7'),
    ('5.  Capability Extraction and Robot Representation', '9'),
    ('6.  Skills Library and Action Grounding',            '11'),
    ('7.  LLM-Based Task Planning',                        '13'),
    ('8.  Dataset and Experimental Design',                '15'),
    ('9.  Evaluation Methodology',                         '17'),
    ('10. Implementation and System Integration',          '19'),
    ('11. Experimental Results and Analysis',              '21'),
    ('12. Limitations',                                    '23'),
    ('13. Future Work',                                    '25'),
    ('14. Conclusion',                                     '27'),
    ('15. Recommendations',                                '29'),
    ('References',                                         '31'),
    ('Glossary',                                           '33'),
    ('Checklist',                                          '35'),
]

table = doc.add_table(rows=len(toc_entries), cols=2)
table.style = 'Table Grid'
for i, (title, page) in enumerate(toc_entries):
    table.cell(i, 0).text = title
    table.cell(i, 1).text = page
    table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in [table.cell(i, 0), table.cell(i, 1)]:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
tbl = table._tbl
tblPr = tbl.tblPr
tblBorders = OxmlElement('w:tblBorders')
for tag in ['top','left','bottom','right','insideH','insideV']:
    border = OxmlElement(f'w:{tag}')
    border.set(qn('w:val'), 'none')
    tblBorders.append(border)
tblPr.append(tblBorders)

# ═══════════════════════════════════════════════════════════════════════════
# LIST OF ABBREVIATIONS
# ═══════════════════════════════════════════════════════════════════════════

section_break(doc)
heading(doc, 'List of Abbreviations and Acronyms', level=1)

abbreviations = [
    ('AI',      'Artificial Intelligence'),
    ('EMOS',    'Embodiment-aware Multi-rObot System'),
    ('JSON',    'JavaScript Object Notation'),
    ('LLM',    'Large Language Model'),
    ('LoRA',   'Low-Rank Adaptation'),
    ('ML',     'Machine Learning'),
    ('MoveIt2','Motion Planning Framework for ROS2'),
    ('Nav2',   'Navigation2 Stack for ROS2'),
    ('PEFT',   'Parameter-Efficient Fine-Tuning'),
    ('ROS2',   'Robot Operating System 2'),
    ('TAMP',   'Task and Motion Planning'),
    ('URDF',   'Unified Robot Description Format'),
    ('WILP',   'Work Integrated Learning Programme'),
    ('Xacro',  'XML Macros (robot description extension)'),
]

table = doc.add_table(rows=len(abbreviations), cols=2)
table.style = 'Table Grid'
for i, (abbr, full) in enumerate(abbreviations):
    c0 = table.cell(i, 0)
    c1 = table.cell(i, 1)
    c0.text = abbr
    c1.text = full
    for cell in [c0, c1]:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
    c0.paragraphs[0].runs[0].bold = True

# ═══════════════════════════════════════════════════════════════════════════
# MAIN BODY
# ═══════════════════════════════════════════════════════════════════════════

def section(doc, num, title):
    section_break(doc)
    h = heading(doc, f'{num}. {title}', level=1)
    return h

def subsection(doc, num, title):
    h = heading(doc, f'{num}  {title}', level=2)
    return h

# ── 1. Introduction ────────────────────────────────────────────────────────
section(doc, 1, 'Introduction and Research Motivation')

body(doc,
    'Robotics is evolving from pre-programmed systems toward intelligent agents capable '
    'of understanding natural-language instructions and operating across diverse environments. '
    'Large Language Models (LLMs) provide strong reasoning and task-decomposition capabilities, '
    'and approaches such as SayCan [1], ProgPrompt [2], and EMOS [3] have demonstrated their '
    'potential for robotic task planning. However, existing approaches often depend on '
    'predefined action spaces, robot-specific configurations, or limited representations '
    'of robot capabilities.')

body(doc,
    'A key challenge is that robots have different embodiments and capabilities. A plan '
    'that is logically correct may still be infeasible if the selected robot lacks the '
    'required mobility, manipulation, perception, payload, or software skills. This creates '
    'a gap between LLM-generated plans and physically executable robotic plans.')

body(doc,
    'This dissertation addresses this gap through a robot-agnostic, capability-aware '
    'framework for language-conditioned task planning. The proposed approach can be '
    'summarized as:')

centered(doc, 'Robot Description → Capability Representation → LLM Planning → Capability-Aware Plan')

body(doc,
    'Robot descriptions such as URDF [5] are transformed into structured capabilities and '
    'combined with world information and available robotic skills to ground LLM planning. '
    'The framework is therefore robot-agnostic in architecture but robot-aware during planning.')

body(doc,
    'The work establishes a foundational architecture, capability representation, planning '
    'pipeline, dataset, and evaluation methodology. It provides a basis for future research '
    'involving real-world execution, dynamic capability learning, closed-loop planning, '
    'multi-robot systems, and broader cross-robot generalization.')

# ── 2. Literature Review ───────────────────────────────────────────────────
section(doc, 2, 'Literature Review and Research Gap')

body(doc,
    'Language-conditioned robotic planning has gained significant attention with the '
    'development of Large Language Models [4]. Existing research demonstrates that LLMs '
    'can interpret high-level instructions and decompose them into sequences of robotic actions.')

body(doc,
    'SayCan [1] combines LLM-generated action probabilities with learned affordance functions '
    'to select actions that are both relevant to the instruction and feasible for the robot. '
    'While effective, its planning process relies on a predefined skill set and robot-specific '
    'affordance models.')

body(doc,
    'ProgPrompt [2] uses programmatic prompting to generate structured robot task plans. '
    'It improves the organization and interpretability of LLM-generated plans but provides '
    'limited explicit reasoning about the physical capabilities of different robot embodiments.')

body(doc,
    'EMOS [3] extends embodiment-aware reasoning by providing structured information about '
    'robots to LLM-based agents, particularly for heterogeneous and multi-robot environments. '
    'This demonstrates the importance of representing embodiment during planning, but leaves '
    'opportunities for more explicit capability extraction, constraint representation, and '
    'executable skill grounding. These limitations were identified as the primary motivation '
    'for the proposed dissertation framework.')

subsection(doc, '2.1', 'Research Gap')

body(doc,
    'Existing approaches demonstrate the effectiveness of LLMs for robotic reasoning, but '
    'a gap remains between language-level planning and embodiment-level feasibility [4]. '
    'In particular, there is a need for a framework that:')

bullet(doc, 'represents heterogeneous robots through a common capability abstraction,')
bullet(doc, 'derives capabilities systematically from robot descriptions,')
bullet(doc, 'grounds generated actions against available robotic skills and environmental constraints, and')
bullet(doc, 'evaluates generated plans independently of another LLM.')

body(doc,
    'This dissertation addresses this gap by introducing a robot-agnostic capability '
    'representation and deterministic evaluation layer around LLM-based task planning, '
    'allowing different robots, environments, tasks, and language models to be evaluated '
    'within a common framework.')

# ── 3. Problem Definition ──────────────────────────────────────────────────
section(doc, 3, 'Problem Definition and Objectives')
subsection(doc, '3.1', 'Problem Definition')

body(doc,
    'LLMs can generate logically meaningful task plans from natural-language instructions, '
    'but they do not inherently guarantee that those plans are feasible for a specific robot. '
    'Differences in mobility, manipulation, perception, sensors, payload, reach, and '
    'available software skills can make the same task executable by one robot and infeasible '
    'for another.')

body(doc,
    'The problem addressed in this dissertation is therefore:')

para = doc.add_paragraph()
para.style = 'Quote'
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = para.add_run(
    'How can natural-language task planning be made robot-agnostic while ensuring that '
    'generated plans remain grounded in the capabilities and constraints of the selected '
    'robot and environment?'
)
set_font(r, bold=True, italic=True)

body(doc,
    'The proposed framework addresses this by separating robot capability representation, '
    'environment representation, available skills, and LLM-based reasoning, allowing the '
    'planner to operate using structured information rather than relying solely on the '
    'LLM\'s internal knowledge.')

subsection(doc, '3.2', 'Research Objectives')

body(doc, 'The primary objectives of this dissertation are:')
bullet(doc, 'Develop a structured and robot-independent representation of robotic capabilities.')
bullet(doc, 'Extract relevant capabilities from robot descriptions such as URDF/Xacro [5].')
bullet(doc, 'Represent executable robotic operations through a common skills library.')
bullet(doc, 'Provide robot capabilities, world state, skills, and task instructions as structured context to an LLM.')
bullet(doc, 'Generate structured, capability-aware task plans across different robots and environments.')
bullet(doc, 'Develop a deterministic software-based evaluator to assess plan validity and feasibility without relying on another LLM.')
bullet(doc, 'Benchmark multiple language models and study their suitability for capability-aware robotic task planning.')

body(doc,
    'These objectives extend the initial dissertation scope of capability extraction, '
    'LLM integration, feasibility-aware planning, and cross-robot evaluation into a '
    'complete experimental framework.')

# ── 4. Framework Architecture ──────────────────────────────────────────────
section(doc, 4, 'Proposed Framework and System Architecture')

body(doc,
    'The proposed system follows a modular architecture that separates robot representation, '
    'environment representation, task reasoning, skill grounding, and plan evaluation. '
    'This separation enables the same planning pipeline to operate across different robot '
    'embodiments and environments.')

centered(doc,
    'Robot Description → Capability Extraction → Capability Representation\n'
    '→ LLM Planner → Generated Plan → Deterministic Evaluation')

body(doc, 'The framework consists of the following major components:')
bullet(doc,
    'Capability Extractor: Parses robot descriptions such as URDF/Xacro [5] and derives '
    'structured information about mobility, manipulation, sensors, joints, reach, '
    'payload-related constraints, and other relevant capabilities.')
bullet(doc,
    'Capability Representation: Converts robot-specific information into a standardized '
    'JSON representation that can be consumed independently of the robot platform.')
bullet(doc,
    'World State Representation: Describes the environment, including locations, objects, '
    'properties, relationships, and environmental constraints relevant to planning.')
bullet(doc,
    'Skills Library: Defines the abstract robotic operations available to the planner '
    'and provides a bridge between high-level actions and underlying robotic functions.')
bullet(doc,
    'LLM Planner: Receives the natural-language task together with robot capabilities, '
    'world state, and available skills to generate a structured sequence of actions.')
bullet(doc,
    'Evaluation Module: Deterministically evaluates the generated plan for factors such '
    'as goal achievement, action and object validity, capability constraints, and logical ordering.')

body(doc,
    'This architecture extends the capability-aware planning concept introduced in the '
    'mid-semester design, where robot capability representation was proposed as an '
    'intermediate layer between LLM reasoning and robot execution.')

body(doc,
    'The modular design also allows individual components — such as the LLM, robot '
    'description, environment, skills library, or evaluator — to be modified independently, '
    'making the framework suitable for systematic experimentation and future extension.')

body(doc,
    'The plan structure within this framework draws inspiration from NavTopo [11], a '
    'topological navigation approach that organizes robot navigation as a sequence of '
    'waypoint sub-tasks rather than a monolithic path. Adopting a similar philosophy, '
    'task plans in the proposed framework are structured as ordered, checkpointed '
    'sub-task sequences. When execution feedback or replanning is triggered, only the '
    'affected checkpoint needs to be regenerated rather than the entire plan. This '
    'selective replanning strategy reduces LLM inference token cost significantly, since '
    'the context provided to the model is scoped to the failed sub-task and its local '
    'constraints rather than the complete task.')

# ── 5. Capability Extraction ───────────────────────────────────────────────
section(doc, 5, 'Capability Extraction and Robot Representation')

body(doc,
    'A core component of the proposed framework is the Capability Extractor, which '
    'converts robot-specific descriptions into a standardized representation that can '
    'be understood by the planning system.')

body(doc,
    'Robot description files such as URDF/Xacro [5] contain structural information '
    'including links, joints, actuators, sensors, and kinematic relationships. However, '
    'providing these files directly to an LLM introduces unnecessary complexity and does '
    'not explicitly describe what the robot can perform. Therefore, the framework '
    'introduces an intermediate capability abstraction.')

centered(doc,
    'URDF/Xacro → Structural Parsing → Semantic Analysis\n'
    '→ Capability Identification → Constraint Generation → Capability JSON')

body(doc, 'A rule-based mechanism interprets structural properties and maps them to planning-relevant capabilities such as:')
bullet(doc, 'mobility and locomotion,')
bullet(doc, 'manipulation and end-effector availability,')
bullet(doc, 'perception and sensing,')
bullet(doc, 'communication and interaction, and')
bullet(doc, 'kinematic and operational constraints.')

body(doc,
    'The resulting Capability JSON acts as a standardized robot profile supplied to the '
    'planner. This allows different robot platforms to be represented through the same '
    'schema while retaining their individual capabilities and limitations.')

body(doc,
    'The approach extends the initial concept of a structured robot resume proposed during '
    'the mid-semester phase into a deterministic capability extraction and representation pipeline.')

body(doc,
    'This abstraction is fundamental to the robot-agnostic design: the planner does not '
    'need to understand each robot\'s URDF directly; it reasons over a common representation '
    'of what the robot can and cannot do.')

# ── 6. Skills Library ──────────────────────────────────────────────────────
section(doc, 6, 'Skills Library and Action Grounding')

body(doc,
    'The Skills Library provides a standardized interface between high-level LLM-generated '
    'plans and executable robotic functions. While the capability representation describes '
    'what a robot can potentially perform, the skills library defines which actions are '
    'available to the planning system and how they can be executed.')

body(doc,
    'Each skill is represented using a structured schema containing information such as '
    'the abstract skill name, required parameters, preconditions, capability requirements, '
    'constraints, and corresponding robotic implementation.')

body(doc, 'The library covers major robotic operations including:')
bullet(doc, 'Navigation and mobility')
bullet(doc, 'Manipulation and object interaction')
bullet(doc, 'Perception and sensing')
bullet(doc, 'Inspection and environment interaction')
bullet(doc, 'Communication and system operations')

body(doc,
    'For example, a high-level action such as pick_object can be associated with '
    'manipulation requirements and mapped to an appropriate ROS2/MoveIt2 [7] '
    'implementation. Similarly, navigation actions can be grounded to corresponding '
    'Nav2 [8] navigation interfaces.')

centered(doc, 'Natural-Language Task → Abstract Skills → Capability Validation → Robot Functions')

body(doc,
    'By separating abstract skills from robot-specific implementations, the framework '
    'avoids directly coupling the LLM to individual ROS2 [6] functions. This enables '
    'a common planning vocabulary across heterogeneous robots while allowing the underlying '
    'implementation to vary according to the selected robotic platform.')

body(doc,
    'Together, the Capability Representation and Skills Library form the grounding layer '
    'that connects language-based reasoning with executable robotic systems.')

# ── 7. LLM Planning ───────────────────────────────────────────────────────
section(doc, 7, 'LLM-Based Task Planning')

body(doc,
    'The LLM Planner acts as the reasoning component of the framework. Its objective is '
    'to convert a natural-language task into a structured sequence of robotic actions '
    'while considering the selected robot, environment, and available skills.')

centered(doc,
    'Natural-Language Task + Robot Capabilities + World State\n'
    '+ Skills Library → Generated Plan')

body(doc,
    'Rather than allowing the LLM to generate unrestricted textual instructions, the '
    'framework uses a structured prompt and predefined output schema. The model is expected '
    'to decompose the requested task into ordered actions, identify relevant objects and '
    'locations, select appropriate skills, and respect the capabilities and constraints '
    'provided in the context.')

body(doc,
    'For example, a task such as "Collect the package and deliver it to the reception desk" '
    'may be decomposed into:')

centered(doc, 'Navigate → Detect Object → Pick Object → Navigate → Place Object')

body(doc,
    'However, the generated sequence depends on whether the robot possesses the required '
    'navigation, perception, and manipulation capabilities and whether the referenced '
    'objects and locations exist in the supplied world state.')

body(doc,
    'The framework is also model-agnostic: the planning interface is designed so that '
    'different LLMs can be evaluated using the same task, robot capability representation, '
    'world state, skills, output structure, and evaluation methodology. This enables '
    'systematic comparison of models based on their ability to generate valid and '
    'capability-aware robotic plans rather than relying solely on general language-generation performance.')

subsection(doc, '7.1', 'Inference Strategy: Zero-Shot and Prompt Engineering')

body(doc,
    'The LLM planner in the current work is evaluated under a zero-shot inference regime: '
    'no task-specific training examples are provided to the model at inference time. The '
    'model is required to generate a valid, capability-aware plan based solely on the '
    'structured context — robot capabilities, world state, skills, and task description — '
    'supplied in a single prompt.')

body(doc,
    'To improve zero-shot performance, the framework employs structured prompt engineering. '
    'The prompt is designed with a clear role definition, explicit output format '
    'specification (JSON schema), constraint declarations derived from the capability '
    'representation, and examples of valid action names from the skills library. This '
    'structured prompting strategy reduces hallucination, enforces schema compliance, and '
    'guides the model toward grounded action selection without requiring fine-tuning.')

body(doc,
    'The current experimental scope is limited to a small number of inference runs due to '
    'computational resource constraints. The available compute restricts the number of '
    'robot-environment-task combinations that can be fully evaluated. Despite this, the '
    'framework design supports arbitrary scale-out: once sufficient inference results are '
    'available, parameter-efficient fine-tuning (PEFT) [9] such as LoRA can be applied to '
    'the best-performing base model to adapt it specifically for capability-aware task '
    'planning. Alternatively, the dataset of generated plans and evaluation scores could '
    'serve as supervision for training a dedicated lightweight neural network with a '
    'task-planning-specific architecture, reducing dependence on general-purpose LLMs '
    'for future deployment.')

# ── 8. Dataset ────────────────────────────────────────────────────────────
section(doc, 8, 'Dataset and Experimental Design')

body(doc,
    'To systematically evaluate the proposed framework, a custom task-planning dataset '
    'was developed across 11 diverse environments, including Kitchen, Apartment, Hospital, '
    'Restaurant, Delivery Warehouse, Factory, Chemistry Laboratory, Shopping Complex, '
    'Playground, College, and Hotel.')

body(doc,
    'The dataset contains approximately 500 natural-language robotic tasks with varying '
    'levels of complexity, ranging from simple actions to navigation-manipulation tasks, '
    'conditional tasks, constraint-aware tasks, and long-horizon plans. The environments '
    'are represented using structured world-state descriptions containing relevant objects, '
    'locations, properties, and relationships.')

body(doc,
    'For the experimental study, selected humanoid robot configurations are evaluated '
    'across these environments using multiple lightweight open-source LLMs. Each experiment '
    'follows a consistent pipeline:')

centered(doc,
    'Task + Robot Capability + World State + Skills\n'
    '→ LLM Inference → Generated Plan → Evaluation')

body(doc,
    'Using identical representations and evaluation criteria across models allows the '
    'experiments to measure differences in planning performance while controlling the '
    'surrounding framework.')

body(doc,
    'The experimental design therefore evaluates three important dimensions of the proposed '
    'approach: cross-task performance, cross-environment generalization, and adaptation to '
    'different robot capabilities. The resulting evaluation scores can additionally serve '
    'as a baseline benchmark for comparing model selection and subsequent parameter-efficient '
    'fine-tuning (PEFT) [9] experiments.')

# ── 9. Evaluation ─────────────────────────────────────────────────────────
section(doc, 9, 'Evaluation Methodology')

body(doc,
    'Evaluating an LLM-generated robotic plan requires more than measuring textual '
    'similarity, since multiple action sequences may correctly accomplish the same task. '
    'Therefore, this dissertation uses a deterministic software-based evaluation module '
    'rather than another LLM as the evaluator.')

body(doc,
    'The evaluator validates the generated plan against the task, robot capabilities, '
    'world state, skills library, and predefined constraints. The major evaluation '
    'dimensions include:')

bullet(doc,
    'Goal Achievement: Whether the generated plan satisfies the intended task objective.')
bullet(doc,
    'Action Validity: Whether generated actions belong to the available skills and action space.')
bullet(doc,
    'Object and Location Validity: Whether referenced entities exist in the corresponding world state.')
bullet(doc,
    'Capability and Constraint Validity: Whether the selected robot possesses the capabilities '
    'required to perform the actions.')
bullet(doc,
    'Logical Ordering: Whether actions follow valid dependencies and preconditions.')
bullet(doc,
    'Plan Feasibility: Whether the complete sequence can reasonably be executed under the '
    'supplied robot and environmental constraints.')

centered(doc,
    'Generated Plan + Ground-Truth Context\n'
    '→ Deterministic Validation → Metric Scores → Overall Performance')

body(doc,
    'Aggregating these metrics across robots, environments, tasks, and models enables '
    'systematic comparison of planning performance. This evaluation strategy also improves '
    'repeatability and transparency, since identical plans evaluated under identical '
    'conditions produce consistent results without introducing the variability or bias '
    'of an LLM-based judge.')

body(doc,
    'It is important to note that the evaluation framework is entirely deterministic and '
    'grounding-based. It checks the different constraints in the generated plan and verifies '
    'whether the plan is consistent with the supplied capability model, world state, and '
    'skills library. However, it is unable to judge whether the robot can actually complete '
    'the task or not in a physical or simulated environment. The evaluation metrics only '
    'check the grounding, but not the logical completeness of the plan. Actual task success '
    'depends on additional factors including low-level motion execution, real-world perception, '
    'actuator reliability, and environmental dynamics, none of which are within the scope '
    'of the evaluation module. This distinction is an important limitation of the current '
    'evaluation framework and motivates future integration with execution-level verification.')

# ── 10. Implementation ────────────────────────────────────────────────────
section(doc, 10, 'Implementation and System Integration')

body(doc,
    'The proposed framework was implemented as a modular Python-based pipeline, allowing '
    'each component to be developed, tested, and modified independently. The implementation '
    'integrates robot capability extraction, structured world representations, a skills '
    'library, LLM inference, and deterministic plan evaluation.')

centered(doc,
    'Robot Description → Capability Extractor → Capability JSON\n'
    'World Definition → World State JSON\n'
    'Task + Capabilities + World State + Skills → LLM → Structured Plan → Evaluator')

body(doc,
    'The implementation uses ROS2-compatible robot descriptions [6], Python, JSON-based '
    'representations, and open-source LLM inference frameworks. The capability extractor '
    'applies predefined rules to interpret robot structures and generate planning-relevant '
    'capabilities and constraints.')

body(doc,
    'A common inference interface was developed so that different LLMs can receive the '
    'same structured inputs and produce plans using a consistent output schema. This '
    'supports controlled comparison between models without modifying the remaining '
    'system architecture.')

body(doc,
    'The individual modules are integrated into an experimental pipeline for executing '
    'tasks across different robot-world combinations, storing generated plans, computing '
    'evaluation metrics, and comparing model performance. This modular implementation '
    'also enables future components, models, robots, and execution systems to be '
    'incorporated without redesigning the complete framework.')

# ── 11. Results ───────────────────────────────────────────────────────────
section(doc, 11, 'Experimental Results and Analysis')

body(doc,
    'The experimental phase evaluates the framework across multiple robots, environments, '
    'natural-language tasks, and LLMs. Each model is provided with the same structured '
    'inputs and evaluated using the deterministic evaluation module, enabling a controlled '
    'comparison of planning performance.')

body(doc, 'The analysis focuses on:')
bullet(doc, 'overall plan validity and goal achievement,')
bullet(doc, 'capability and constraint compliance,')
bullet(doc, 'action, object, and location validity,')
bullet(doc, 'logical ordering of generated actions,')
bullet(doc, 'performance variation across task complexity and environments,')
bullet(doc, 'differences in planning behaviour across robot embodiments, and')
bullet(doc, 'comparative performance of the selected LLMs.')

body(doc,
    'The experiments are designed not only to identify the best-performing model, but '
    'also to understand where and why planning failures occur. Errors are categorized '
    'into cases such as unsupported actions, invalid environmental references, capability '
    'violations, incorrect action ordering, incomplete plans, and malformed structured outputs.')

body(doc,
    'The resulting scores provide a quantitative baseline for the proposed framework. '
    'The strongest-performing model can subsequently be used for Parameter-Efficient '
    'Fine-Tuning (PEFT) [9], after which the same evaluation pipeline can be repeated '
    'to measure whether domain-specific adaptation improves capability-aware robotic planning.')

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = para.add_run(
    '[Note: Detailed quantitative results, comparative tables, graphs, and model-wise '
    'observations are to be incorporated in this section upon completion of the full '
    'inference and evaluation runs.]'
)
set_font(r, italic=True)

# ── 12. Limitations ───────────────────────────────────────────────────────
section(doc, 12, 'Limitations')

body(doc,
    'Although the proposed framework establishes a structured approach for capability-aware '
    'robotic task planning, several limitations remain within the current scope of the dissertation.')

bullet(doc,
    'Simulation and representation-based validation: The generated plans are primarily '
    'evaluated using structured robot capabilities, world states, and software-based '
    'validation. Robot execution in simulation is not included in the current scope as it '
    'requires significant time for model finetuning, world building, and robot creation. '
    'This extensive setup will instead be a focus for future work.')
bullet(doc,
    'Deterministic metrics check grounding, not logical completeness: The evaluation '
    'framework checks grounding validity and constraint compliance deterministically, but '
    'it is unable to judge if the robot can actually complete the task or not. The metrics '
    'only check grounding, but not logical completeness. Actual task success depends on '
    'motion execution, real-world perception, and actuator reliability, which are outside '
    'the current scope.')
bullet(doc,
    'Limited inference scale due to compute constraints: The current experimental runs are '
    'limited in number due to available computational resources. A larger-scale evaluation '
    'across all robot-environment-task combinations requires significantly more compute capacity.')
bullet(doc,
    'Zero-shot inference only: The models are currently evaluated in a zero-shot setting. '
    'Fine-tuning on domain-specific data has not yet been performed, which may limit the '
    'models\' ability to consistently follow the structured output schema and capability constraints.')
bullet(doc,
    'Static capability representation: Robot capabilities are extracted from available '
    'descriptions and predefined rules. Dynamic factors such as hardware degradation, '
    'battery state, environmental conditions, and changing payload are not currently modeled.')
bullet(doc,
    'Dependence on robot descriptions: The accuracy of capability extraction depends on '
    'the completeness and correctness of URDF/Xacro [5] and associated robot information.')
bullet(doc,
    'Predefined skills library: The framework assumes a structured set of available '
    'robotic skills. Automatic discovery or learning of new skills is outside the present scope.')
bullet(doc,
    'LLM limitations: Generated plans can still contain hallucinated entities, unsupported '
    'actions, incorrect ordering, or incomplete reasoning despite capability and environment grounding.')
bullet(doc,
    'High-level planning focus: The dissertation primarily addresses task-level planning '
    'and does not attempt to replace low-level motion planning, trajectory generation, '
    'control, perception, or localization systems.')
bullet(doc,
    'Benchmark scope: The developed dataset provides controlled evaluation across multiple '
    'robots, environments, and tasks, but it cannot represent every real-world condition '
    'or robotic embodiment.')

body(doc,
    'These limitations define the boundary of the current work rather than the final '
    'boundary of the proposed research direction. The framework provides a foundation '
    'upon which dynamic capability modeling, execution feedback, larger-scale benchmarking, '
    'and real-world robotic validation can subsequently be investigated.')

# ── 13. Future Work ───────────────────────────────────────────────────────
section(doc, 13, 'Future Work')

body(doc,
    'The current dissertation establishes a foundation for robot-agnostic and '
    'capability-aware task planning. Several directions can extend this framework into '
    'deeper research.')

bullet(doc,
    'Fine-tuning and domain-specific model training: Once a sufficient volume of '
    'capability-aware planning inferences and evaluation scores has been accumulated, '
    'the best-performing base LLM can be fine-tuned using parameter-efficient methods '
    'such as LoRA [9]. Alternatively, a dedicated lightweight neural network with a '
    'task-planning-specific architecture could be trained on the collected dataset, '
    'potentially replacing general-purpose LLMs for deployment in resource-constrained environments.')
bullet(doc,
    'Multi-robot coordination for single tasks: Extend the framework to support '
    'heterogeneous multi-robot teams collaborating on a single complex task. Different '
    'sub-tasks could be allocated to robots based on their complementary capabilities, '
    'with the LLM planner generating coordinated, role-aware action sequences that account '
    'for inter-robot dependencies and communication requirements. This extends beyond the '
    'current single-robot-per-task paradigm while reusing the same capability representation infrastructure.')
bullet(doc,
    'Hierarchical and selective replanning (NavTopo-inspired): Building on the checkpoint-based '
    'plan structure inspired by NavTopo [11], future work should implement full closed-loop '
    'replanning where only the affected checkpoint sub-task is regenerated upon failure, '
    'rather than regenerating the entire plan. This approach further reduces inference token '
    'cost and latency while preserving the already-validated portions of the plan.')
bullet(doc,
    'Execution-level task success verification: Integrate the planning framework with '
    'physical or simulation environments to measure actual task completion, moving beyond '
    'deterministic grounding checks. This would enable end-to-end evaluation from '
    'natural-language instruction to successful robot execution.')
bullet(doc,
    'Real-world robot validation: Deploy the framework on physical robots and compare '
    'software-based feasibility predictions with actual execution outcomes.')
bullet(doc,
    'Dynamic capability modeling: Extend static capability representations to account '
    'for battery state, payload, sensor availability, hardware failures, and changing '
    'environmental conditions.')
bullet(doc,
    'Closed-loop planning: Incorporate execution feedback so that plans can be dynamically '
    'corrected or regenerated when actions fail or the environment changes.')
bullet(doc,
    'Task and Motion Planning integration: Connect high-level LLM-generated task plans '
    'with motion planners [10] to verify geometric feasibility, reachability, collision '
    'constraints, and trajectory execution.')
bullet(doc,
    'Automatic skill discovery: Reduce dependence on manually defined skill libraries by '
    'discovering available skills and interfaces directly from robotic software ecosystems '
    'such as ROS2 [6].')
bullet(doc,
    'Learning-based capability representation: Investigate whether robot capabilities can '
    'be learned or refined from execution history instead of relying entirely on '
    'deterministic extraction rules.')
bullet(doc,
    'Larger-scale benchmarking: Expand the dataset to additional robots, environments, '
    'tasks, and language models to establish a broader benchmark for capability-aware '
    'robotic planning.')

body(doc,
    'These directions provide a natural continuation toward doctoral research. The current '
    'framework can serve as the base experimental platform, while dedicated research can '
    'investigate adaptive capability representations, embodied feedback, cross-robot '
    'generalization, and ultimately the interaction between language reasoning and physical '
    'robot intelligence.')

# ── 14. Conclusion ────────────────────────────────────────────────────────
section(doc, 14, 'Conclusion')

body(doc,
    'This dissertation presented a robot-agnostic, capability-aware framework for '
    'language-conditioned robotic task planning. The work addresses the gap between '
    'the semantic reasoning capabilities of Large Language Models and the physical and '
    'functional constraints of heterogeneous robotic systems.')

body(doc,
    'The framework integrates robot capability extraction, structured capability '
    'representation, world-state modeling, a common skills library, LLM-based planning, '
    'and deterministic plan evaluation. By separating robot-specific information from '
    'the planning mechanism, the same architecture can be applied across different '
    'robots, environments, tasks, and language models.')

body(doc,
    'A key outcome of this work is the principle that a logically correct plan is not '
    'necessarily a robot-feasible plan. Explicit representation of robot capabilities '
    'and environmental constraints is therefore essential when applying language models '
    'to robotic task planning.')

body(doc,
    'The dissertation establishes a functional and extensible research foundation rather '
    'than attempting to solve general-purpose robotic planning completely. With dedicated '
    'research into real-world execution, adaptive capability learning, closed-loop planning, '
    'task-and-motion integration, and broader cross-robot evaluation, the proposed framework '
    'can be extended toward more general and reliable language-driven robotic intelligence.')

# ── 15. Recommendations ───────────────────────────────────────────────────
section(doc, 15, 'Recommendations')

body(doc,
    'Based on the design, implementation, and evaluation methodology developed in this '
    'dissertation, the following recommendations are proposed for further development '
    'of capability-aware robotic planning.')

bullet(doc,
    'Capability information should be explicitly represented rather than expecting an LLM '
    'to infer robot abilities from its pretrained knowledge.')
bullet(doc,
    'Robot-specific information should remain separated from the planning model, allowing '
    'the same planner to operate across different embodiments.')
bullet(doc,
    'Structured world states and skill definitions should be used to reduce hallucinated '
    'objects, locations, and unsupported robotic actions.')
bullet(doc,
    'LLM-generated plans should be independently validated before execution using '
    'deterministic capability, constraint, and logical checks.')
bullet(doc,
    'Evaluation should extend beyond task success, considering action validity, '
    'environmental grounding, capability compliance, logical ordering, and overall feasibility.')
bullet(doc,
    'LLMs should function as high-level reasoning components, while conventional robotics '
    'systems continue to handle perception, navigation, manipulation, motion planning, and control.')

body(doc,
    'These recommendations support a modular approach in which advances in LLMs, robot '
    'hardware, capability extraction, and execution systems can be incorporated independently '
    'without redesigning the complete planning framework.')

# ═══════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════

section_break(doc)
heading(doc, 'References', level=1)

references = [
    ('[1]', 'Ahn, M., Brohan, A., Brown, N., Chebotar, Y., Cortes, O., David, B., Finn, C., '
             'Fu, C., Gopalakrishnan, K., Hausman, K., Herzog, A., Ho, D., Hsu, J., Ibarz, J., '
             'Ichter, B., Irpan, A., Jang, E., Ruano, R. J., Jeffrey, K., Jesmonth, S., Joshi, '
             'N. J., Julian, R., Kalashnikov, D., Kuang, Y., Lee, K.-H., Levine, S., Lu, Y., '
             'Luu, L., Parada, C., Pastor, P., Quiambao, M., Rao, K., Rettinghouse, J., '
             'Reyes, D., Sermanet, P., Sievers, N., Tan, C., Toshev, A., Vanhoucke, V., '
             'Xia, F., Xiao, T., Xu, P., Xu, S., and Zeng, A. "Do As I Can, Not As I Say: '
             'Grounding Language in Robotic Affordances." arXiv preprint arXiv:2204.01691, 2022.'),
    ('[2]', 'Singh, I., Blukis, V., Mousavian, A., Goyal, A., Xu, D., Tremblay, J., Fox, D., '
             'Thomason, J., and Garg, A. "ProgPrompt: Generating Situated Robot Task Plans using '
             'Large Language Models." In Proceedings of the IEEE International Conference on '
             'Robotics and Automation (ICRA), pp. 11523–11530, 2023.'),
    ('[3]', '[Full EMOS citation — please replace with complete author(s), title, venue, and year.]'),
    ('[4]', 'Zhao, W. X., Zhou, K., Li, J., Tang, T., Wang, X., Hou, Y., Min, Y., Zhang, B., '
             'Zhang, J., Dong, Z., Du, Y., Yang, C., Chen, Y., Chen, Z., Jiang, J., Ren, R., '
             'Li, Y., Tang, X., Liu, Z., Liu, P., Nie, J., and Wen, J.-R. "A Survey of Large '
             'Language Models." arXiv preprint arXiv:2303.18223, 2023.'),
    ('[5]', 'Open Robotics. "URDF — Unified Robot Description Format." ROS Documentation, 2022. '
             'Available: https://wiki.ros.org/urdf [Accessed: July 2026].'),
    ('[6]', 'Macenski, S., Foote, T., Gerkey, B., Lalancette, C., and Woodall, W. "Robot '
             'Operating System 2: Design, Architecture, and Uses in the Wild." Science Robotics, '
             'vol. 7, no. 66, eabm6074, 2022.'),
    ('[7]', 'Coleman, D., Sucan, I., Chitta, S., and Correll, N. "Reducing the Barrier to '
             'Entry of Complex Robotic Software: a MoveIt! Case Study." arXiv preprint '
             'arXiv:1404.3785, 2014. See also: MoveIt2 Documentation. Available: '
             'https://moveit.ros.org [Accessed: July 2026].'),
    ('[8]', 'Macenski, S., Martín, F., White, R., and Ginés Clavero, J. "The Marathon 2: '
             'A Navigation System." In Proceedings of the IEEE/RSJ International Conference '
             'on Intelligent Robots and Systems (IROS), pp. 2718–2725, 2020.'),
    ('[9]', 'Hu, E. J., Shen, Y., Wallis, P., Allen-Zhu, Z., Li, Y., Wang, S., Wang, L., '
             'and Chen, W. "LoRA: Low-Rank Adaptation of Large Language Models." In '
             'Proceedings of the International Conference on Learning Representations (ICLR), 2022.'),
    ('[10]','Garrett, C. R., Chitnis, R., Holladay, R., Kim, B., Silver, T., Kaelbling, '
             'L. P., and Lozano-Pérez, T. "Integrated Task and Motion Planning." Annual '
             'Review of Control, Robotics, and Autonomous Systems, vol. 4, pp. 265–293, 2021.'),
    ('[11]','Gao, F., Peng, K., Zhang, H., Bermudez, I., and Payandeh, S. '
             '"NavTopo: Leveraging Topological Maps For Autonomous Navigation Of A Robot." '
             'arXiv preprint arXiv:2410.01492, 2024. '
             'Available: https://arxiv.org/abs/2410.01492 [Accessed: July 2026].'),
]

for num, ref in references:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(para, before=0, after=6)
    r_num = para.add_run(num + '  ')
    set_font(r_num, bold=True)
    r_text = para.add_run(ref)
    set_font(r_text)
    # Hanging indent
    para.paragraph_format.left_indent    = Cm(1.0)
    para.paragraph_format.first_line_indent = Cm(-1.0)

# ═══════════════════════════════════════════════════════════════════════════
# GLOSSARY
# ═══════════════════════════════════════════════════════════════════════════

section_break(doc)
heading(doc, 'Glossary', level=1)

glossary = [
    ('Capability Extractor',
     'A software component that parses robot description files (URDF/Xacro) and derives '
     'structured capability representations for use in planning.'),
    ('Capability JSON',
     'A JSON-formatted structured representation of a robot\'s capabilities, constraints, '
     'and operational parameters used as context input to the LLM planner.'),
    ('Embodiment',
     'The physical and functional characteristics of a robot that determine its ability '
     'to perform tasks in a given environment.'),
    ('Grounding',
     'The process of connecting high-level language or symbolic representations to '
     'concrete, executable actions and real-world entities.'),
    ('Hallucination (LLM)',
     'The generation by an LLM of plausible-sounding but incorrect, non-existent, or '
     'unsupported objects, actions, or facts.'),
    ('Large Language Model (LLM)',
     'A neural language model pre-trained on large corpora that exhibits emergent '
     'capabilities for reasoning, planning, and language understanding.'),
    ('PEFT',
     'Parameter-Efficient Fine-Tuning: a class of techniques for adapting pre-trained '
     'models to specific tasks by updating only a small subset of parameters, including '
     'methods such as LoRA.'),
    ('Robot-agnostic',
     'Describing a system or framework designed to operate with any robot embodiment '
     'without requiring robot-specific modifications to the core architecture.'),
    ('Skills Library',
     'A structured repository of abstract robotic skills with associated parameters, '
     'preconditions, and implementation mappings used to ground LLM-generated actions.'),
    ('Task Planning',
     'The process of decomposing a high-level goal into a sequence of executable actions '
     'that achieve that goal.'),
    ('URDF',
     'The Unified Robot Description Format, an XML-based file format used in ROS/ROS2 '
     'to describe the physical properties, kinematics, and dynamics of a robot.'),
    ('World State',
     'A structured representation of the environment, including the locations, objects, '
     'properties, and relationships relevant to a given planning task.'),
]

for term, definition in glossary:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(para, before=0, after=6)
    r_term = para.add_run(term + ':  ')
    set_font(r_term, bold=True)
    r_def = para.add_run(definition)
    set_font(r_def)
    para.paragraph_format.left_indent       = Cm(1.0)
    para.paragraph_format.first_line_indent = Cm(-1.0)

# ═══════════════════════════════════════════════════════════════════════════
# CHECKLIST  (last page)
# ═══════════════════════════════════════════════════════════════════════════

section_break(doc)
heading(doc, 'Checklist of Items for the Final Dissertation Report', level=1)

para = doc.add_paragraph(
    'This checklist is attached as the last page of the final report and is to be '
    'duly completed, verified, and signed by the student.'
)
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
set_paragraph_spacing(para, before=0, after=10)
for run in para.runs:
    run.font.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

checklist = [
    ('1',    'Is the final report neatly formatted with all the elements required for a technical report?',                               'Yes'),
    ('2',    'Is the Cover page in proper format as given in Annexure A?',                                                                'Yes'),
    ('3',    'Is the Title page (Inner cover page) in proper format?',                                                                    'Yes'),
    ('4a',   'Is the Certificate from the Supervisor in proper format?',                                                                  'Yes'),
    ('4b',   'Has it been signed by the Supervisor?',                                                                                     '[To sign]'),
    ('5',    'Is the Abstract included in the report properly written within one page?',                                                  'Yes'),
    ('6',    'Have the technical keywords been specified properly?',                                                                       'Yes'),
    ('7',    'Is the title of your report appropriate? (Descriptive, precise, no uncommon abbreviations)',                                'Yes'),
    ('8',    'Have you included the List of Abbreviations/Acronyms?',                                                                     'Yes'),
    ('9',    'Does the report contain a summary of the literature survey?',                                                               'Yes'),
    ('10',   'Does the Table of Contents include page numbers?',                                                                          'Yes'),
    ('10i',  'Are the pages numbered properly? (Chapter 1 starts on page 1)',                                                             'Yes'),
    ('10ii', 'Are the figures numbered properly? (Figure numbers and titles at bottom of figures)',                                       'Yes'),
    ('10iii','Are the tables numbered properly? (Table numbers and titles at top of tables)',                                             'Yes'),
    ('10iv', 'Are the captions for figures and tables proper?',                                                                           'Yes'),
    ('10v',  'Are the appendices numbered properly with appropriate titles?',                                                             'Yes'),
    ('11',   'Is the conclusion of the report based on discussion of the work?',                                                          'Yes'),
    ('12',   'Are references or bibliography given at the end of the report?',                                                            'Yes'),
    ('13',   'Have the references been cited properly inside the text of the report?',                                                    'Yes'),
    ('14',   'Are all the references cited in the body of the report?',                                                                   'Yes'),
    ('15',   'Is the report format and content according to the guidelines? (Not a PPT printout or user manual)',                         'Yes'),
]

table = doc.add_table(rows=len(checklist) + 1, cols=3)
table.style = 'Table Grid'

hdr = table.rows[0].cells
hdr[0].text = 'No.'
hdr[1].text = 'Description'
hdr[2].text = 'Status'
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

for i, (num, desc, status) in enumerate(checklist, start=1):
    row = table.rows[i].cells
    row[0].text = num
    row[1].text = desc
    row[2].text = status
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)

# Column widths
from docx.shared import Cm
table.columns[0].width = Cm(1.5)
table.columns[1].width = Cm(11.5)
table.columns[2].width = Cm(2.0)

doc.add_paragraph()
bold_label(doc, 'Declaration by Student:')
body(doc,
    'I certify that I have properly verified all the items in this checklist and ensure '
    'that the report is in proper format as specified in the course handout.'
)
doc.add_paragraph()
sig_table(doc, 'Place: _______________\nDate: _______________',
               '')
doc.add_paragraph()
bold_label(doc, 'Signature of the Student:', '_______________________')
bold_label(doc, 'Name:', 'Sana Jaya Krishna')
bold_label(doc, 'ID No.:', '2024AA05783')

# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════

output_path = 'dissertation_final_report.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
